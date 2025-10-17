from docx import Document
from pathlib import Path

def export_docx(clean_text: str, marked_html: str, docx_path: Path):
    doc = Document()
    doc.add_heading('Reintext (bereinigt)', level=1)
    for para in clean_text.split('\n\n'):
        doc.add_paragraph(para)
    doc.add_page_break()
    doc.add_heading('Mit Fehler-Markierung (HTML)', level=1)
    doc.add_paragraph('Hinweis: Diese Sektion enthält HTML-Markup vereinfacht als Text.')
    doc.add_paragraph(marked_html)
    doc.save(docx_path)
